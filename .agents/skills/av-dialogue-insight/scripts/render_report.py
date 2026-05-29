#!/usr/bin/env python3
"""Render a consolidated audio/video dialogue analysis into a structured report.

Input is a JSON analysis document (produced by the agent from one or more
`analyze_media` calls, optionally merged with `merge_chunks.py`). Output is
deterministic Markdown with: 概要, 事件时间线, 说话人画像, 情感时间线, 关键触发点,
总结. A `.docx` is also emitted when `--docx` is passed and python-docx is
installed.

Analysis shape (all sections optional except `media`):

    {
      "media": "meeting.mp4",
      "duration_seconds": 95.0,
      "summary": "……",
      "degraded": false,
      "degraded_note": "",
      "events": [{"time": "00:05", "title": "...", "detail": "..."}],
      "speakers": [{"id": "S1", "label": "主持人", "talk_ratio": 0.45, "profile": "..."}],
      "emotion_timeline": [{"time": "00:05", "speaker": "S1", "emotion": "neutral", "valence": 0.0, "note": "..."}],
      "key_triggers": [{"time": "01:12", "description": "...", "evidence": "..."}]
    }

Usage:
    render_report.py <analysis.json> <output_base> [--docx]
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def md_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        cells = [str(cell).replace("|", "\\|").replace("\n", " ") for cell in row]
        lines.append("| " + " | ".join(cells) + " |")
    return lines


def render_markdown(analysis: dict) -> str:
    media = str(analysis.get("media") or "media")
    lines: list[str] = [f"# 音视频对话分析报告：{media}", ""]

    duration = analysis.get("duration_seconds")
    if duration is not None:
        lines.append(f"时长：{float(duration):.1f}s")
    if analysis.get("degraded"):
        lines.append(f"> ⚠️ 降级处理：{analysis.get('degraded_note') or '部分能力不可用，结果可能不完整。'}")
    if duration is not None or analysis.get("degraded"):
        lines.append("")

    if analysis.get("summary"):
        lines += ["## 概要", str(analysis["summary"]).strip(), ""]

    events = analysis.get("events") or []
    if events:
        lines += ["## 事件时间线"]
        lines += md_table(
            ["时间", "事件", "说明"],
            [[e.get("time", ""), e.get("title", ""), e.get("detail", "")] for e in events],
        )
        lines.append("")

    speakers = analysis.get("speakers") or []
    if speakers:
        lines += ["## 说话人画像"]
        lines += md_table(
            ["编号", "标签", "话语占比", "画像"],
            [
                [
                    s.get("id", ""),
                    s.get("label", ""),
                    f"{float(s['talk_ratio']) * 100:.0f}%" if s.get("talk_ratio") is not None else "",
                    s.get("profile", ""),
                ]
                for s in speakers
            ],
        )
        lines.append("")

    emotions = analysis.get("emotion_timeline") or []
    if emotions:
        lines += ["## 情感时间线"]
        lines += md_table(
            ["时间", "说话人", "情感", "效价", "依据"],
            [
                [
                    em.get("time", ""),
                    em.get("speaker", ""),
                    em.get("emotion", ""),
                    f"{float(em['valence']):+.1f}" if em.get("valence") is not None else "",
                    em.get("note", ""),
                ]
                for em in emotions
            ],
        )
        lines.append("")

    triggers = analysis.get("key_triggers") or []
    if triggers:
        lines += ["## 关键触发点"]
        for t in triggers:
            lines.append(f"- **{t.get('time', '')}** {t.get('description', '')}".rstrip())
            if t.get("evidence"):
                lines.append(f"  - 依据：{t['evidence']}")
        lines.append("")

    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines) + "\n"


def render_docx(analysis: dict, output_path: Path) -> bool:
    try:
        from docx import Document  # type: ignore
    except Exception:
        return False

    document = Document()
    document.add_heading(f"音视频对话分析报告：{analysis.get('media', 'media')}", level=0)
    if analysis.get("summary"):
        document.add_heading("概要", level=1)
        document.add_paragraph(str(analysis["summary"]))

    def add_table(title: str, headers: list[str], rows: list[list[str]]) -> None:
        if not rows:
            return
        document.add_heading(title, level=1)
        table = document.add_table(rows=1, cols=len(headers))
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row in rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

    add_table("事件时间线", ["时间", "事件", "说明"], [[e.get("time", ""), e.get("title", ""), e.get("detail", "")] for e in analysis.get("events", [])])
    add_table("说话人画像", ["编号", "标签", "画像"], [[s.get("id", ""), s.get("label", ""), s.get("profile", "")] for s in analysis.get("speakers", [])])
    add_table("情感时间线", ["时间", "说话人", "情感", "依据"], [[em.get("time", ""), em.get("speaker", ""), em.get("emotion", ""), em.get("note", "")] for em in analysis.get("emotion_timeline", [])])

    if analysis.get("key_triggers"):
        document.add_heading("关键触发点", level=1)
        for t in analysis["key_triggers"]:
            document.add_paragraph(f"{t.get('time', '')} {t.get('description', '')}", style="List Bullet")

    document.save(str(output_path))
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description="Render an A/V dialogue analysis into a report.")
    parser.add_argument("analysis", help="Path to the consolidated analysis JSON.")
    parser.add_argument("output_base", help="Output path without extension; .md/.docx is appended.")
    parser.add_argument("--docx", action="store_true", help="Also emit a .docx (requires python-docx).")
    args = parser.parse_args()

    try:
        analysis = json.loads(Path(args.analysis).read_text(encoding="utf-8"))
    except FileNotFoundError:
        print(f"analysis not found: {args.analysis}", file=sys.stderr)
        return 1
    except json.JSONDecodeError as error:
        print(f"invalid analysis JSON: {error}", file=sys.stderr)
        return 1

    if not isinstance(analysis, dict):
        print("analysis must be a JSON object", file=sys.stderr)
        return 1

    output_base = Path(args.output_base)
    output_base.parent.mkdir(parents=True, exist_ok=True)

    markdown_path = output_base.with_suffix(".md")
    markdown_path.write_text(render_markdown(analysis), encoding="utf-8")
    print(f"wrote {markdown_path}")

    if args.docx:
        docx_path = output_base.with_suffix(".docx")
        if render_docx(analysis, docx_path):
            print(f"wrote {docx_path}")
        else:
            print("python-docx not available; skipped .docx", file=sys.stderr)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
