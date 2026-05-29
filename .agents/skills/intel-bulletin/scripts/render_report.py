#!/usr/bin/env python3
"""Render a structured bulletin spec into a 公文-format report.

Input is a JSON spec describing the bulletin; output is deterministic Markdown
(always) and, when `python-docx` is importable and `--docx` is passed, a `.docx`
alongside it.

Spec shape (all fields optional except `title` and `sections`):

    {
      "title": "关于XX情况的情报通报",
      "doc_number": "情报〔2026〕第008号",
      "classification": "内部",
      "recipient": "各相关部门",
      "summary": "一段概要性说明……",
      "sections": [
        {"heading": "基本情况", "body": "……"},
        {"heading": "分析研判", "body": "……"}
      ],
      "conclusion": "综上……",
      "issuer": "情报分析组",
      "date": "2026-05-29"
    }

Usage:
    render_report.py <spec.json> <output_base> [--docx]

`<output_base>` is a path without extension; `.md` (and optionally `.docx`) is
appended.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

CN_DIGITS = "零一二三四五六七八九"


def cn_ordinal(n: int) -> str:
    """Render 1..99 as Chinese numerals (一、二、…十、十一、…)."""
    if n <= 0:
        return str(n)
    if n < 10:
        return CN_DIGITS[n]
    if n == 10:
        return "十"
    if n < 20:
        return "十" + CN_DIGITS[n % 10]
    tens, ones = divmod(n, 10)
    return CN_DIGITS[tens] + "十" + (CN_DIGITS[ones] if ones else "")


def render_markdown(spec: dict) -> str:
    title = str(spec.get("title") or "情报通报").strip()
    lines: list[str] = [f"# {title}", ""]

    doc_number = spec.get("doc_number")
    classification = spec.get("classification")
    if classification:
        lines.append(f"密级：{str(classification).strip()}")
    if doc_number:
        lines.append(f"编号：{str(doc_number).strip()}")
    if classification or doc_number:
        lines.append("")

    recipient = spec.get("recipient")
    if recipient:
        lines.append(f"{str(recipient).strip()}：")
        lines.append("")

    summary = spec.get("summary")
    if summary:
        lines.append(f"【概要】{str(summary).strip()}")
        lines.append("")

    sections = spec.get("sections") or []
    for index, section in enumerate(sections, start=1):
        heading = str(section.get("heading") or "").strip()
        body = str(section.get("body") or "").strip()
        lines.append(f"{cn_ordinal(index)}、{heading}")
        if body:
            lines.append(body)
        lines.append("")

    conclusion = spec.get("conclusion")
    if conclusion:
        lines.append(f"【结论】{str(conclusion).strip()}")
        lines.append("")

    issuer = spec.get("issuer")
    date = spec.get("date")
    if issuer or date:
        lines.append("———")
        if issuer:
            lines.append(f"落款：{str(issuer).strip()}")
        if date:
            lines.append(f"日期：{str(date).strip()}")

    # Collapse trailing blank lines into a single terminal newline.
    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines) + "\n"


def render_docx(spec: dict, output_path: Path) -> bool:
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except Exception:
        return False

    document = Document()
    title = document.add_heading(str(spec.get("title") or "情报通报").strip(), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = []
    if spec.get("classification"):
        meta.append(f"密级：{spec['classification']}")
    if spec.get("doc_number"):
        meta.append(f"编号：{spec['doc_number']}")
    if meta:
        document.add_paragraph("    ".join(meta))
    if spec.get("recipient"):
        document.add_paragraph(f"{spec['recipient']}：")
    if spec.get("summary"):
        document.add_paragraph(f"【概要】{spec['summary']}")

    for index, section in enumerate(spec.get("sections") or [], start=1):
        document.add_heading(f"{cn_ordinal(index)}、{section.get('heading', '')}", level=1)
        if section.get("body"):
            document.add_paragraph(str(section["body"]))

    if spec.get("conclusion"):
        document.add_paragraph(f"【结论】{spec['conclusion']}")
    if spec.get("issuer") or spec.get("date"):
        document.add_paragraph("")
        if spec.get("issuer"):
            p = document.add_paragraph(f"落款：{spec['issuer']}")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if spec.get("date"):
            p = document.add_paragraph(f"日期：{spec['date']}")
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.save(str(output_path))
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a bulletin spec into a 公文-format report.")
    parser.add_argument("spec", help="Path to the bulletin spec JSON file.")
    parser.add_argument("output_base", help="Output path without extension; .md/.docx is appended.")
    parser.add_argument("--docx", action="store_true", help="Also emit a .docx (requires python-docx).")
    args = parser.parse_args()

    spec_path = Path(args.spec)
    try:
        spec = json.loads(spec_path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        print(f"spec not found: {spec_path}", file=sys.stderr)
        return 1
    except json.JSONDecodeError as error:
        print(f"invalid spec JSON: {error}", file=sys.stderr)
        return 1

    if not isinstance(spec, dict):
        print("spec must be a JSON object", file=sys.stderr)
        return 1

    output_base = Path(args.output_base)
    output_base.parent.mkdir(parents=True, exist_ok=True)

    markdown_path = output_base.with_suffix(".md")
    markdown_path.write_text(render_markdown(spec), encoding="utf-8")
    print(f"wrote {markdown_path}")

    if args.docx:
        docx_path = output_base.with_suffix(".docx")
        if render_docx(spec, docx_path):
            print(f"wrote {docx_path}")
        else:
            print("python-docx not available; skipped .docx (install python-docx to enable)", file=sys.stderr)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
